from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from pathlib import Path

ROOT = Path('/mnt/data/picorv32_assembler_project')
OUT = ROOT / 'PicoRV32_RV32I_Assembler_Raporu_Tam_Uyumlu.docx'

opcodes = [
    ('add','R','33','0','00','rd, rs1, rs2'), ('sub','R','33','0','20','rd, rs1, rs2'),
    ('sll','R','33','1','00','rd, rs1, rs2'), ('slt','R','33','2','00','rd, rs1, rs2'),
    ('sltu','R','33','3','00','rd, rs1, rs2'), ('xor','R','33','4','00','rd, rs1, rs2'),
    ('srl','R','33','5','00','rd, rs1, rs2'), ('sra','R','33','5','20','rd, rs1, rs2'),
    ('or','R','33','6','00','rd, rs1, rs2'), ('and','R','33','7','00','rd, rs1, rs2'),
    ('addi','I','13','0','00','rd, rs1, imm'), ('slti','I','13','2','00','rd, rs1, imm'),
    ('sltiu','I','13','3','00','rd, rs1, imm'), ('xori','I','13','4','00','rd, rs1, imm'),
    ('ori','I','13','6','00','rd, rs1, imm'), ('andi','I','13','7','00','rd, rs1, imm'),
    ('slli','I','13','1','00','rd, rs1, shamt'), ('srli','I','13','5','00','rd, rs1, shamt'),
    ('srai','I','13','5','20','rd, rs1, shamt'), ('lb','I','03','0','00','rd, imm(rs1)'),
    ('lh','I','03','1','00','rd, imm(rs1)'), ('lw','I','03','2','00','rd, imm(rs1)'),
    ('lbu','I','03','4','00','rd, imm(rs1)'), ('lhu','I','03','5','00','rd, imm(rs1)'),
    ('jalr','I','67','0','00','rd, imm(rs1)'), ('sb','S','23','0','00','rs2, imm(rs1)'),
    ('sh','S','23','1','00','rs2, imm(rs1)'), ('sw','S','23','2','00','rs2, imm(rs1)'),
    ('beq','B','63','0','00','rs1, rs2, label'), ('bne','B','63','1','00','rs1, rs2, label'),
    ('blt','B','63','4','00','rs1, rs2, label'), ('bge','B','63','5','00','rs1, rs2, label'),
    ('bltu','B','63','6','00','rs1, rs2, label'), ('bgeu','B','63','7','00','rs1, rs2, label'),
    ('lui','U','37','-','-','rd, imm20'), ('auipc','U','17','-','-','rd, imm20'),
    ('jal','J','6F','-','-','rd, label'), ('ecall','SYS','73','-','-','-'), ('ebreak','SYS','73','-','-','-')
]

report = Document()
section = report.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

styles = report.styles
normal = styles['Normal']
normal.font.name = 'Courier New'
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
normal.font.size = Pt(10)

for style_name in ['Title', 'Subtitle', 'Heading 1', 'Heading 2', 'Heading 3']:
    st = styles[style_name]
    st.font.name = 'Courier New'
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
    st.font.size = Pt(10)
    st.font.bold = True
styles['Title'].font.size = Pt(12)
styles['Subtitle'].font.size = Pt(10)


def font_run(run, size=10, bold=False, italic=False):
    run.font.name = 'Courier New'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def shade_cell(cell, fill='D9E2F3'):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, align='left', size=9.5):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    font_run(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(text, level=1):
    report.add_heading(text, level=level)


def add_para(text='', bold=False, center=False, italic=False):
    p = report.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    font_run(run, size=10, bold=bold, italic=italic)
    return p


def add_bullets(items):
    for item in items:
        p = report.add_paragraph(style='Normal')
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run('- ' + item)
        font_run(run, size=10)


def add_code(text):
    p = report.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    run = p.add_run(text)
    font_run(run, size=9)


def add_table(headers, rows, col_widths=None, fill='D9EAD3'):
    table = report.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, head in enumerate(headers):
        set_cell_text(hdr[i], head, bold=True, align='center')
        shade_cell(hdr[i], fill)
        if col_widths:
            hdr[i].width = col_widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            sval = str(val)
            align = 'center' if len(sval) <= 12 and '\n' not in sval else 'left'
            set_cell_text(cells[i], sval, align=align)
            if col_widths:
                cells[i].width = col_widths[i]
    report.add_paragraph('')


def add_box(title, content):
    table = report.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.2)
    hdr = table.cell(0, 0)
    body = table.cell(1, 0)
    shade_cell(hdr, 'B7C9E2')
    shade_cell(body, 'F7F7F7')
    set_cell_text(hdr, title, bold=True, align='center', size=10)
    p = body.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(content)
    font_run(run, size=9)
    report.add_paragraph('')


# Cover
add_para('PicoRV32 Islemcisi Icin RV32I Alt Kumesi Assembler Tasarimi ve Object Kod Uretimi', bold=True, center=True)
add_para('C Dili ile Iki Gecisli (Two-Pass) Assembler Gerceklestirimi', center=True)
add_para('Cikti Bicimi: Intel HEX   |   Hedef: PicoRV32 uyumlu RV32I alt kumesi', center=True)
add_para('Hazirlayan: [Ogrenci adi - numara]', center=True)
add_para('Ders: Sistem Programlama', center=True)
add_para('Rapor yazim bicimi: Courier New, 10 pt', center=True)
report.add_paragraph('')
add_table(
    ['Oge', 'Bu rapordaki durum'],
    [
        ['Yazilim dili', 'C99'],
        ['Assembler turu', 'Two-pass'],
        ['Desteklenen komut siniflari', 'R, I, S, B, U, J, ecall, ebreak'],
        ['Desteklenen direktifler', '.text, .data, .word, .byte, .org, .end'],
        ['Teslim artefaktlari', 'Kaynak kod, README, testler, HEX, listing, rapor'],
    ],
    [Cm(4.2), Cm(10.0)],
    fill='E2F0D9'
)
report.add_page_break()

# 1
add_heading('1. Giris', 1)
add_para('Bu projenin amaci, PicoRV islemcisine ait RV32I komut kumesinin bir alt kumesi icin tam calisan bir assembler tasarlamak ve bu tasarimi raporlamaktir. Proje tanimina gore gelistirilen yazilimin opcode table, symbol table, parser, makine kodu uretici ve .data/.text/.word/.byte/.org/.end direktiflerinin en az bir kismini desteklemesi beklenmektedir. Rapor tarafinda ise giris, literatur, assembler mimarisi, veri yapilari, opcode/symbol table tasarimi, assembler algoritmasi, akis diyagrami, test senaryolari, karmasiklik analizi ve sonuc bolumlerinin bulunmasi istenmistir.')
add_para('Bu raporda sunulan gerceklestirim C99 ile yazilmistir ve iki gecisli assembler mantigi kullanir. Pass 1 asamasinda adres atama ve sembol toplama islemleri yapilir; Pass 2 asamasinda ise komutlar 32 bitlik RISC-V makine koduna cevrilir, veri direktifleri byte duzeyinde uretilir ve Intel HEX cikisi yazilir. Iki gecisli tercih, ileri referanslarin guvenilir sekilde cozulmesini sagladigi icin secilmistir.')
add_bullets([
    'Hedef platform: PicoRV32 ile uyumlu RV32I taban komut alt kumesi.',
    'Uygulama dili: C99.',
    'Nesne kod cikisi: Intel HEX.',
    'Dogrulama artefaktlari: .hex ve .lst listing dosyalari.',
    'Tasarim ilkeleri: modulerlik, gozlenebilirlik, genisletilebilirlik, tekrarlanabilir test.'
])

# 2
add_heading('2. Literatur Arastirmasi', 1)
add_para('Assembler tasarimi literaturde tipik olarak OPTAB, SYMTAB, LOCCTR ve two-pass isleme modeli etrafinda aciklanir. Ders notlarinda da temel assembler fonksiyonlari; mnemonic->opcode cevirisi, sembolik etiketlere adres atama, veri sabitlerini object koda cevirme ve Pass 1/Pass 2 ayrimi olarak ifade edilmistir. Bu projede mimari ve veri yapilari, bu klasik yaklasimin RV32I mimarisi icin yeniden uyarlanmasi ile kurulmustur.')
add_para('RISC-V tarafinda resmi unprivileged architecture dokumani, sabit 32 bit komut uzunlugu, temel instruction formatlari ve immediate alanlarinin dagitimi icin ana kaynaktir. Bu sabitlik assembler tasarimini kolaylastirmistir; cunku komut sinifi belirlendikten sonra alanlarin bit konumlari sabittir. RV32I 2.1 durumu resmi RISC-V ratified specification kutuphanesinde listelenmektedir. PicoRV32 dokumani ise cekirdegin RV32I dahil farkli konfigurasiyonlarda kullanilabildigini gosterdiginden destek kapsamimizi acik bir alt kume olarak tanimlamamiz gerekmiştir.')
add_para('Klasik assembler literaturu sadece teori degil, veri yapisi seciminin pratik etkisini de vurgular. Modern Compiler Design bolumunde assemblerin en kritik gorevleri; sembol cozumleme, nesne kod uretimi ve object format yazimi olarak ele alinmistir. Buna ek olarak, mnemonic arama performansinin assembler tasariminda onemli oldugunu inceleyen guncel calismalar, OPTAB ve SYMTAB icin uygun tablo organizasyonunun derleme verimini etkiledigini gostermektedir.')
add_para('Egitim odakli calismalarda ise assembler, derleyici ve islemci tasariminin birlikte ele alinmasi, ogrencinin mimari ile sistem yazilimi arasindaki baglantiyi daha iyi kavramasini saglamaktadir. Bu nedenle bu raporda sadece kod verilmemis; blok diyagram, akis diyagrami, test programlari, listing ciktisi ve rubrik odakli PÇ basliklari da eklenmistir.')
add_table(
    ['Kaynak', 'Bu projedeki kullanimi'],
    [
        ['RISC-V Unprivileged ISA', 'Komut formatlari, opcode/funct alanlari, immediate yerlestirmesi'],
        ['PicoRV32 teknik dokumani', 'Hedef cekirdegin RV32I uyumlu alt kumesinin sinirlanmasi'],
        ['Assembler ders notlari', 'OPTAB, SYMTAB, LOCCTR, two-pass algoritma omurgasi'],
        ['Modern Compiler Design', 'Assembler modullerinin ayrilmasi ve object code uretim mantigi'],
        ['Guncel akademik makaleler', 'Veri yapisi ve egitsel tasarim gerekceleri'],
    ],
    [Cm(5.0), Cm(9.2)]
)

# 3
add_heading('3. Assembler Mimarisi', 1)
add_para('Yazilim mimarisi yedi temel module ayrilmistir: Parser, LineVector, Opcode Table, Symbol Table, Pass 1, Pass 2 ve Intel HEX Writer. Parser kaynak satiri label, mnemonic, operand ve yorum alanlarina ayirir. Pass 1, adres atama ve sembol tablosu uretir. Pass 2 ise ayni ara temsil uzerinden komutlari kodlar ve bellek imgesini olusturur. Son olarak HEX writer, kullanilan adresleri kayitlara donusturur.')
add_box('Blok Diyagram', r'''
+--------------------+
| Assembly Source .s |
+---------+----------+
          |
          v
+--------------------+
| Parser / Tokenizer |
+---------+----------+
          |
          v
+--------------------+        +--------------------+
| Pass 1             |------->| Symbol Table       |
| LOCCTR + Labels    |        | label -> address   |
+---------+----------+        +--------------------+
          |
          v
+--------------------+        +--------------------+
| Pass 2             |------->| Opcode Table       |
| Encoder            |        | RV32I metadata     |
+---------+----------+        +--------------------+
          |
          v
+--------------------+
| Memory Image       |
+---------+----------+
          |
          v
+--------------------+
| Intel HEX Writer   |
+--------------------+
''')
add_table(
    ['Modul', 'Girdi', 'Cikti', 'Temel sorumluluk'],
    [
        ['Parser', 'Kaynak satir', 'ParsedLine', 'Label, mnemonic, operand ayristirma'],
        ['Pass 1', 'ParsedLine dizisi', 'Adres bilgisi + SYMTAB', 'LOCCTR ilerletme ve label kaydi'],
        ['Pass 2', 'ParsedLine + SYMTAB', 'Makine kodu byte dizileri', 'Instruction/directive encoding'],
        ['Memory Image', 'Uretilen byte lar', 'Adreslenebilir bellek gorunumu', 'HEX yazimi oncesi duz adres uzayi'],
        ['HEX Writer', 'Memory Image', '.hex dosyasi', 'Intel HEX record uretimi'],
    ],
    [Cm(3.0), Cm(3.0), Cm(4.0), Cm(4.2)]
)
add_para('Bu ayrim sayesinde her modulin bagimsiz test edilmesi mumkun olmustur. Ornegin parser hatalari listing cikisindan, branch/jump ofset hatalari ise listing ve test assembly dosyalarindan izlenebilmektedir. Mimari ayni zamanda gelecekte pseudo-instruction veya yeni direktif eklenmesini de kolaylastirir.')

# 4
add_heading('4. Kullanilan Veri Yapilari', 1)
add_para('Projede secilen veri yapilari dogrudan rubrikte istenen veri yapisi tasarimi maddesine cevap verecek sekilde belirlenmistir. Kaynak satirlar dinamik bir LineVector icinde tutulur. Bu tasarim, dosyanin ikinci kez fiziksel olarak okunmasini gerektirmez; Pass 1 sonunda uretilen adresler ve parser sonucu, ayni satir nesnesi uzerinde Pass 2 ye aktarilir.')
add_para('Symbol Table acik adreslemeli hash tablo olarak gerceklestirilmistir. Her kayit label ismi, adresi, section bilgisi, absolute/relative bayragi ve tanimli-mi bilgisini tutar. Opcode Table ise sabit oldugu icin dinamik arama yapisina degil, salt-okunur statik diziye yerlestirilmistir. Memory Image yapisinda data[] gercek byte degerlerini, used[] ise ilgili adrese veri yazilip yazilmadigini gosterir. Bu secim, Intel HEX yazma asamasinda bos adres alanlarini atlamayi saglar.')
add_table(
    ['Yapi', 'Amac', 'Gerceklestirim', 'Avantaj'],
    [
        ['LineVector', 'Ara temsil', 'Dinamik dizi', 'Pass1 ve Pass2 ayni parsed satirlari kullanir'],
        ['Symbol Table', 'Label -> adres', 'Acik adreslemeli hash tablo', 'Ortalama O(1) ekleme ve arama'],
        ['Opcode Table', 'Mnemonic -> metadata', 'Sabit statik dizi', 'Dusuk bellek, dusuk bakim maliyeti'],
        ['Memory Image', 'Adres -> byte', 'Paralel data/used dizileri', 'HEX uretiminde yalniz kullanilan adresler yazilir'],
        ['ErrorList', 'Toplu hata raporu', 'Dizi', 'Ilk hatada dusmek yerine tum baglam korunur'],
    ],
    [Cm(3.0), Cm(3.5), Cm(4.0), Cm(4.0)]
)

# 5
add_heading('5. Opcode Table Tasarimi', 1)
add_para('Opcode tablosu her mnemonic icin gerekli minimum bilgiyi tutar: instruction formati, opcode, funct3, funct7 ve operand kalibi. RISC-V tarafinda register alanlari tum ilgili formatlarda ayni bit konumlarina yerlestigi icin Pass 2, once mnemonic icin tablo girdisini bulur, sonra operandlari sayisal alanlara cevirir ve bitleri yerlestirir.')
add_para('Bu projede desteklenen alt kume; R-type aritmetik ve mantiksal komutlar, I-type immediate ve load komutlari, S-type store komutlari, B-type kosullu dallanmalar, U-type ust immediate komutlari, J-type jal komutu ve ecall/ebreak sistem komutlaridir. FENCE ve CSR ailesi kapsam disi tutulmustur; cunku proje amaci temel assembler omurgasini gostermektir ve PicoRV32 odakli egitsel alt kumede zorunlu degildir.')
add_table(['Mnemonic','Fmt','Opcode','funct3','funct7','Operand kalibi'], opcodes, [Cm(2.0), Cm(1.0), Cm(1.7), Cm(1.5), Cm(1.7), Cm(6.3)])

# 6
add_heading('6. Symbol Table Tasarimi', 1)
add_para('SYMTAB, assemblerin ileri referans problemini cozmesinde ana veri yapisidir. Pass 1 asamasinda gorulen her label, o andaki LOCCTR degeriyle tabloya kaydedilir. Ayni etiket ikinci kez tanimlanirsa duplicate label hatasi uretillir. Pass 2 asamasinda branch, jal, .word ve .byte ifadeleri de dahil olmak uzere sembolik operandlar bu tablo uzerinden cozulur.')
add_para('Acik adreslemeli hash tablo secimi, dusuk bagimlilik ve ortalama O(1) davranis sagladigi icin yapilmistir. FNV-1a benzeri bir hash ile benzer on ekler tasiyan etiket isimlerinde dahi kabul edilebilir dagilim elde edilmistir. Yuk faktoru arttiginda rehash yapilarak carpismalarin zincirleme buyumesi onlenir.')
add_bullets([
    'Anahtar: label ismi.',
    'Deger: adres, section, absolute/relative bilgisi, tanimli-mi bayragi.',
    'Hata durumu: duplicate label ve undefined symbol raporlanir.',
    'Kullanim yerleri: branch hedefleri, jal/jalr hedefleri, .org/.word/.byte ifadeleri.',
])

# 7
add_heading('7. Assembler Algoritmasi', 1)
add_para('Algoritma iki gecislidir. Parser once butun satirlari ayrıştırir. Pass 1, LOCCTR sayaci uzerinden adres atar ve label lari SYMTAB icine yazar. Pass 2 ise mnemonic, register ve immediate alanlarini bitlere yerlestirerek 32 bitlik makine kodunu uretir; veri direktifleri icin ise byte dizileri olusturulur. Son asamada Memory Image Intel HEX kayitlarina cevrilir.')
add_box('Akis Diyagrami', r'''
BASLA
  |
  v
Kaynak dosyayi oku
  |
  v
Satirlari parser ile ayrıştır
  |
  v
PASS 1
  - label varsa SYMTAB'a ekle
  - .org varsa LOCCTR'yi guncelle
  - .word/.byte boyut hesapla
  - komut ise +4 byte ilerle
  |
  v
Hata var mi?
  |---- Evet ---> hata raporu yaz ve dur
  |
  Hayir
  v
PASS 2
  - mnemonic icin OPTAB kaydi bul
  - operandlari sayisal alana donustur
  - instruction/directive bytes uret
  - Memory Image'a yaz
  |
  v
Intel HEX ve listing uret
  |
  v
BITIR
''')
add_heading('7.1 Assembler tasariminda kullanilan butun argumanlarin etkisi', 2)
add_para('Proje duyurusunda ozellikle assembler tasariminda kullanilan butun argumanlarin etkisinin bilinmesi ve raporda yer almasi istenmistir. Bu nedenle desteklenen operand ve direktif argumanlari asagida ayri bir tabloda ozetlenmistir.')
add_table(
    ['Arguman/alan', 'Kullanildigi yer', 'Assemblera etkisi'],
    [
        ['rd', 'R/I/U/J', 'Hedef register numarasi olarak bit [11:7] alanina yazilir'],
        ['rs1', 'R/I/S/B', 'Kaynak register 1 olarak bit [19:15] alanina yazilir'],
        ['rs2', 'R/S/B', 'Kaynak register 2 olarak bit [24:20] alanina yazilir'],
        ['imm12', 'I-type ve load/jalr', 'Isaretli 12 bit immediate olarak bit [31:20] alanina yazilir'],
        ['imm12 (split)', 'S-type', 'Ust kisim [31:25], alt kisim [11:7] alanina dagitilir'],
        ['label / branch hedefi', 'B-type ve jal', 'PC goreli ofset hesaplanir; B/J tiplerinde daginik bit alanlarina yeniden yerlestirilir'],
        ['.org adresi', 'Direktif', 'LOCCTR dogrudan yeni adrese ayarlanir'],
        ['.word ifadesi', 'Direktif', 'Her operand icin 4 byte veri uretir ve little-endian yazar'],
        ['.byte ifadesi', 'Direktif', 'Her operand icin 1 byte veri uretir'],
        ['.text / .data', 'Direktif', 'Mantiksal bolum degisimi belirtir; mevcut uygulamada tek duz adres uzayi surdurulur'],
        ['.end', 'Direktif', 'Derleme akisini sonlandirir'],
    ],
    [Cm(3.3), Cm(3.2), Cm(8.5)]
)
add_para('Bu tasarimda ifadeler; sayi, karakter literal i, label, label+imm ve label-imm bicimlerini destekler. Boylece temel veri yerlesimi ve sembolik adresleme senaryolari kapsanmis olur. Branch ve jump komutlarinda ofsetin hizalama ve bit alanina sigma kosulu ayrica dogrulanir.')

# 8
add_heading('8. Test Senaryolari ve Sonuclar', 1)
add_para('Dogrulama icin uc farkli test assembly programi kullanilmistir. Testler bilincli olarak farkli komut formatlarini, veri direktiflerini ve kontrol akislarini kapsayacak sekilde secilmistir. Her test icin hem Intel HEX hem de listing (.lst) dosyasi uretilmistir.')
add_table(
    ['Test', 'Amac', 'Kapsam', 'Uretilen dosyalar'],
    [
        ['Test-1', 'Dongu ve geri dallanma', 'addi, blt, ecall', 'test1.hex, test1.lst'],
        ['Test-2', 'Kod-veri yerlesimi', '.data, .text, .org, lw, sw, .word', 'test2.hex, test2.lst'],
        ['Test-3', 'Alt program cagrisi', 'jal, jalr, add, ecall', 'test3.hex, test3.lst'],
    ],
    [Cm(2.2), Cm(3.5), Cm(6.0), Cm(3.5)]
)
add_para('Test-1, negatif branch ofseti ile geri atlayan bir dongu icerir. Bu test B-type immediate alaninin dogru dagitilmasini ve PC goreli hesaplamanin saglikli calistigini gostermektedir.')
add_code('00000000    93 02 00 00    addi t0, zero, 0')
add_code('00000004    13 03 50 00    addi t1, zero, 5')
add_code('00000008    93 82 12 00    addi t0, t0, 1')
add_code('0000000C    E3 CE 62 FE    blt  t0, t1, loop')
add_code('00000010    73 00 00 00    ecall')
add_para('Test-2, .org ve .data/.text gecisleri ile bellek yerlesimini dogrular. Bu test sayesinde assemblerin yalniz komut degil, veri direktiflerini de dogru adreslere yerlestirdigi ve HEX writer in bos adres bloklarini atlayabildigi gozlenmistir.')
add_code('00000008    03 A3 02 00    lw   t1, 0(t0)')
add_code('00000010    23 A2 62 00    sw   t1, 4(t0)')
add_code('00000100    0A 00 00 00 14 00 00 00    .word 10, 20')
add_code('00000108    73 00 10 00    ebreak')
add_para('Test-3, jal ve jalr kullanan alt program akisini sinar. Boylece hem J-type immediate kodlamasi hem de donus adresinin ra registerina yazilmasi listing duzeyinde izlenebilir hale gelir.')
add_code('00000008    EF 00 80 00    jal  ra, add_func')
add_code('00000010    33 05 B5 00    add  a0, a0, a1')
add_code('00000014    67 80 00 00    jalr zero, 0(ra)')
add_table(
    ['Dogrulama olcutu', 'Sonuc'],
    [
        ['Parser tum test dosyalarini hatasiz ayrıştırdi', 'Evet'],
        ['Pass 1 label adreslerini dogru atadi', 'Evet'],
        ['Pass 2 RV32I kodlamalarini uretti', 'Evet'],
        ['Intel HEX cikisi olustu', 'Evet'],
        ['Listing dosyalari ile elde edilen byte lar kontrol edildi', 'Evet'],
    ],
    [Cm(10.5), Cm(3.5)],
    fill='FFF2CC'
)

# 9
add_heading('9. Algoritma Karmasikligi Analizi (Time / Space Complexity)', 1)
add_para('N kaynak satir sayisini, S sembol sayisini, B uretilen toplam byte sayisini gostersin. OPTAB sabit boyutlu oldugundan arama maliyeti pratikte O(1) kabul edilmistir. SYMTAB acik adreslemeli hash tablo oldugu icin ekleme ve arama ortalama O(1), en kotu durumda O(S) olabilir.')
add_table(
    ['Asama', 'Ortalama zaman', 'En kotu zaman', 'Alan'],
    [
        ['Parser', 'O(N)', 'O(N)', 'O(N)'],
        ['Pass 1', 'O(N)', 'O(N*S) patolojik carpismada', 'O(S)'],
        ['Pass 2', 'O(N + B)', 'O(N*S + B)', 'O(B + S)'],
        ['Intel HEX yazimi', 'O(B)', 'O(B)', 'O(1) ek alan'],
    ],
    [Cm(3.0), Cm(4.0), Cm(5.4), Cm(3.0)]
)
add_para('Dolayisiyla toplam ortalama zaman karmasikligi O(N + B), toplam alan karmasikligi ise O(N + S + B) olarak ifade edilebilir. Ara temsil, sembol tablosu ve bellek imgesi ayni anda tutuldugu icin alan maliyeti bu uc ana yapinin toplami ile belirlenir. Sabit 32 bit instruction boyu, Pass 1 de komut boyu kararini sabit zamana indirerek algoritmayi sadeleştirmistir.')

# 10
add_heading('10. Program Ciktilari - PÇ6, PÇ7, PÇ8, PÇ12, PÇ13, PÇ17', 1)
add_para('Proje duyurusunda PÇ6, PÇ7, PÇ8, PÇ12, PÇ13 ve PÇ17 kapsaminda yapilan calismalarin ayri basliklar halinde verilmesi istendigi icin bu bolum rubrik ile uyumlu sekilde yeniden duzenlenmistir. PÇ17 tanimi kurum sisteminden dogrulanmalidir; burada PÇ17, acik kaynak ekosistemini izleme ve yeni teknolojileri ogrenme ekseninde ele alinmistir.')
add_heading('10.1 PÇ6 - Problemin analizi, mimari, veri yapisi ve algoritma tasarimi', 2)
add_para('PÇ6 kapsaminda en belirgin katki, assembler probleminin alt gorevlere ayrilmasi ve her gorevin uygun veri yapisi ile eslenmesidir. OPTAB, SYMTAB, LOCCTR, parser ve encoder ayrimi sayesinde problem sadece "assembly kodunu cevir" seviyesinde birakilmamis; instruction formatlari, label cozumleme, veri yerlesimi ve object format uretimi olarak sistematik bicimde analiz edilmistir. Mimari tasarim, veri yapisi secimi ve Big-O analizi bu PÇ icin dogrudan kanittir.')
add_heading('10.2 PÇ7 - Yazilim gerceklestirme ve test/dogrulama', 2)
add_para('PÇ7 kapsaminda proje, calisan C99 kaynak kodu, derlenebilir Makefile yapisi, gercek test assembly dosyalari ve uretilebilen HEX/listing dosyalari ile somutlasmistir. Yalniz teorik tasarim yapilmamis; parser, pass1, pass2, opcode table, symbol table ve hex writer modulleri gerceklenmistir. Uc test senaryosu ile farkli instruction aileleri ve direktifler dogrulanmistir.')
add_heading('10.3 PÇ8 - Muhendislik cozumunun etkileri', 2)
add_para('Bu assembler, acik bir ISA olan RISC-V uzerinde calistigi icin egitim ve arastirma ortamlarinda tekrar uretilebilir bir arac olarak deger tasir. PicoRV32 gibi FPGA odakli cekirdeklerle uyumlu nesne kod uretimi, dusuk maliyetli deneysel sistemlerde kullanilabilecek bir arac ortaya koyar. Toplumsal ve teknolojik etki acisindan bakildiginda; acik standartlarin, tekrarlanabilir deneylerin ve tasarim seffafliginin desteklenmesi olumlu bir muhendislik sonucudur.')
add_heading('10.4 PÇ12 - Takim calismasi ve sunum hazirligi', 2)
add_para('Bu calisma bireysel yurutulmus olsa dahi PÇ12 rubrigiyle uyumlu olacak sekilde is paketlerine ayrilmistir: mimari tasarim, parser/encoder gelistirme, test, raporlama ve sunum hazirligi. Grup projesine donusturulmesi halinde ayni tablo kisi bazli is bolumune dogrudan cevrilebilir. Sunum acisindan ise kodun canli calistirilmasi, testlerin listing dosyalariyla gosterilmesi ve rapordaki blok/akis diyagramlarinin kullanilmasi hedeflenmistir.')
add_table(
    ['Is paketi', 'Uretilen artefakt'],
    [
        ['Mimari tasarim', 'Blok diyagram, moduller, veri akis tanimi'],
        ['Gerceklestirim', 'C99 kaynak kodu ve Makefile'],
        ['Dogrulama', '3 test programi, .hex ve .lst ciktilari'],
        ['Dokumantasyon', 'Akademik rapor, README'],
        ['Sunum hazirligi', '20 dk teknik akisa uygun icerik plani'],
    ],
    [Cm(4.5), Cm(9.0)]
)
add_heading('10.5 PÇ13 - Literatur arastirmasi ve rapor kalitesi', 2)
add_para('PÇ13 kapsaminda rapor, yalnizca bir sonuc dokumani degil; literaturle iliskilendirilmis teknik bir aciklama metni olarak kurgulanmistir. Resmi RISC-V dokumanlari, PicoRV32 teknik dokumani, assembler tasarimi ile ilgili akademik kaynaklar ve ders notlari birlikte kullanilarak mimari kararlar gerekcelendirilmistir. Rapor duzeninde baslik hiyerarsisi, tablolar, diyagramlar ve test listingleri kullanilarak akademik yazim niteligine yaklasilmistir.')
add_heading('10.6 PÇ17 - Acik kaynak ekosistemi ve surekli ogrenme', 2)
add_para('PÇ17 tanimi bolumden bolume degisebildigi icin bu raporda dikkatli bir yorum benimsenmistir. Projede resmi RISC-V spesifikasyonlari, PicoRV32 acik kaynak deposu ve literaturden elde edilen bilgiler kullanilarak guncel ve ogrenmeye acik bir gelistirme sureci izlenmistir. Acik ISA ve acik cekirdek etrafinda arac gelistirmek; yeni mimarileri takip etme, teknik dokumani okuyup uygulamaya aktarma ve surekli kendini guncelleme becerisini gostermektedir.')
add_heading('10.7 Degerlendirme kriterleri ile uyum ozeti', 2)
add_table(
    ['Kriter', 'Rapor/artefakt karsiligi'],
    [
        ['Problemin Analizi (PÇ6)', 'Giris, mimari, veri yapilari, algoritma ve karmasiklik bolumleri'],
        ['Literatur Arastirmasi (PÇ13)', 'Bolum 2 ve kaynakca'],
        ['Assembler Mimarisi (PÇ6)', 'Bolum 3 ve blok diyagram'],
        ['Veri Yapilari (PÇ6)', 'Bolum 4, 5 ve 6'],
        ['Algoritma Tasarimi (PÇ6)', 'Bolum 7 ve akis diyagrami'],
        ['Yazilim Gerceklestirme (PÇ7)', 'Kaynak kodlar, Makefile, README'],
        ['Test ve Dogrulama (PÇ7)', 'Bolum 8, .hex ve .lst dosyalari'],
        ['Muhendislik Etkisi (PÇ8)', 'Bolum 10.3'],
        ['Takim calismasi / Sunum (PÇ12)', 'Bolum 10.4 ve Ek-A'],
        ['Rapor Kalitesi (PÇ13)', 'Courier New 10 pt duzende akademik rapor'],
        ['PÇ17', 'Bolum 10.6'],
    ],
    [Cm(5.0), Cm(8.5)],
    fill='DDEBF7'
)

# 11
add_heading('11. Sonuc ve Degerlendirme', 1)
add_para('Bu projede PicoRV32 odakli RV32I alt kumesi icin calisan, iki gecisli ve Intel HEX cikisi ureten bir assembler gelistirilmistir. Yazilim; opcode table, symbol table, parser, pass1, pass2 ve hex writer modullerinden olusan temiz bir mimariye sahiptir. Testler ile branch, load/store, veri yerlesimi, alt program cagrisi ve direktif isleme mekanizmalari dogrulanmistir.')
add_para('Projenin guclu yonleri; moduler C99 tasarimi, hash tabanli sembol tablosu, listing dosyalari ile gozlenebilirlik, resmi RISC-V dokumanlariyla uyumlu instruction encoding ve rubrik odakli raporlama yapisidir. Sinirlari ise pseudo-instruction destegi, relocation kayitlari, CSR/FENCE ailesi, makro sistemi ve daha zengin expression parser in henuz eklenmemis olmasidir. Bu maddeler, gelecekteki gelistirme adimlari olarak dogal bir yol haritasi sunmaktadir.')
add_para('Genel degerlendirme olarak, proje duyurusunda istenen mimari tasarim, veri yapilari, algoritma, yazilim gerceklestirme, test, rapor ve PÇ basliklari buyuk olcude karsilanmis; eksik kalan kisimlar ise bu revizyonla PÇ17 ve rubrik uyumlu PÇ alt basliklari eklenerek tamamlanmistir.')

# 12
add_heading('12. Kaynakca', 1)
refs = [
    '[1] RISC-V International, The RISC-V Instruction Set Manual, Volume I: Unprivileged Architecture, official release, version 20260120, 2026.',
    '[2] A. Waterman, K. Asanovic (eds.), The RISC-V Instruction Set Manual, Volume I: Unprivileged Architecture, version 20250508, 2025.',
    '[3] K. Asanovic and D. A. Patterson, Instruction Sets Should Be Free: The Case For RISC-V, UCB/EECS-2014-146, University of California, Berkeley, 2014.',
    '[4] D. Grune, C. J. H. Jacobs, K. G. Langendoen, H. E. Bal and K. Langendoen, Modern Compiler Design, 2nd ed., Springer, Chapter 8: Assemblers, Disassemblers, Linkers, and Loaders, 2012.',
    '[5] H. Oztekin, A. Lazzem and I. Pehlivan, Using FPGA-based content-addressable memory for mnemonics instruction searching in assembler design, The Journal of Supercomputing, vol. 79, no. 15, pp. 17386-17418, 2023.',
    '[6] K. Nakano and Y. Ito, Processor, Assembler, and Compiler Design Education Using an FPGA, 14th IEEE International Conference on Parallel and Distributed Systems (ICPADS), pp. 723-728, 2008.',
    '[7] YosysHQ, PicoRV32 technical documentation and repository, GitHub, accessed 2026.',
    '[8] Ders Notlari, Temel assembler tasarimi, komut formatlari, makineden bagimsiz assembler ozellikleri ve relocation konulari, Sistem Programlama dersi materyali.',
]
for ref in refs:
    add_para(ref)

# Appendix
add_heading('Ek-A. 20 Dakikalik Sunum Icin Onerilen Akis', 1)
add_table(
    ['Sure', 'Baslik', 'Sunum icerigi'],
    [
        ['2 dk', 'Problemin tanimi', 'Neden assembler, neden PicoRV32/RV32I, proje kapsam maddeleri'],
        ['4 dk', 'Mimari tasarim', 'Blok diyagram, moduller ve veri akisinin aciklanmasi'],
        ['4 dk', 'Veri yapilari', 'OPTAB, SYMTAB, Memory Image, neden hash tablo secildigi'],
        ['4 dk', 'Algoritma', 'Pass 1 / Pass 2 akisi ve operand etkileri'],
        ['4 dk', 'Canli gosterim', 'Assemblerin calistirilmasi, test1/test2/test3 ve listing/HEX ciktilari'],
        ['2 dk', 'Sonuc', 'Sinirlar, gelistirme onerileri, sorular'],
    ],
    [Cm(2.0), Cm(3.5), Cm(8.0)],
    fill='EADCF4'
)

report.save(OUT)
print(OUT)
