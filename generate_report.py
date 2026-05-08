from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from pathlib import Path

ROOT = Path('/mnt/data/picorv32_assembler_project')
OUT = ROOT / 'PicoRV32_RV32I_Assembler_Raporu.docx'

opcodes = [
    ('add','R','33','0','00','rd, rs1, rs2'), ('sub','R','33','0','20','rd, rs1, rs2'),
    ('sll','R','33','1','00','rd, rs1, rs2'), ('slt','R','33','2','00','rd, rs1, rs2'),
    ('sltu','R','33','3','00','rd, rs1, rs2'), ('xor','R','33','4','00','rd, rs1, rs2'),
    ('srl','R','33','5','00','rd, rs1, rs2'), ('sra','R','33','5','20','rd, rs1, rs2'),
    ('or','R','33','6','00','rd, rs1, rs2'), ('and','R','33','7','00','rd, rs1, rs2'),
    ('addi','I','13','0','00','rd, rs1, imm'), ('slti','I','13','2','00','rd, rs1, imm'),
    ('sltiu','I','13','3','00','rd, rs1, imm'), ('xori','I','13','4','00','rd, rs1, imm'),
    ('ori','I','13','6','00','rd, rs1, imm'), ('andi','I','13','7','00','rd, rs1, imm'),
    ('slli','I','13','1','00','rd, rs1, shamt'), ('srli','I','13','5','00','rd, rs1, shamt'),
    ('srai','I','13','5','20','rd, rs1, shamt'), ('lb','I','03','0','00','rd, imm(rs1)'),
    ('lh','I','03','1','00','rd, imm(rs1)'), ('lw','I','03','2','00','rd, imm(rs1)'),
    ('lbu','I','03','4','00','rd, imm(rs1)'), ('lhu','I','03','5','00','rd, imm(rs1)'),
    ('jalr','I','67','0','00','rd, imm(rs1)'), ('sb','S','23','0','00','rs2, imm(rs1)'),
    ('sh','S','23','1','00','rs2, imm(rs1)'), ('sw','S','23','2','00','rs2, imm(rs1)'),
    ('beq','B','63','0','00','rs1, rs2, label'), ('bne','B','63','1','00','rs1, rs2, label'),
    ('blt','B','63','4','00','rs1, rs2, label'), ('bge','B','63','5','00','rs1, rs2, label'),
    ('bltu','B','63','6','00','rs1, rs2, label'), ('bgeu','B','63','7','00','rs1, rs2, label'),
    ('lui','U','37','-','-','rd, imm20'), ('auipc','U','17','-','-','rd, imm20'),
    ('jal','J','6F','-','-','rd, label'), ('ecall','SYS','73','-','-','-'), ('ebreak','SYS','73','-','-','-')
]

report = Document()
section = report.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

styles = report.styles
normal = styles['Normal']
normal.font.name = 'Courier New'
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
normal.font.size = Pt(10)

for style_name in ['Title', 'Subtitle', 'Heading 1', 'Heading 2', 'Heading 3']:
    st = styles[style_name]
    st.font.name = 'Courier New'
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
    st.font.size = Pt(10)
    st.font.bold = True

styles['Title'].font.size = Pt(12)
styles['Subtitle'].font.size = Pt(10)


def shade_cell(cell, fill='D9E2F3'):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, align='left'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'right': WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
    run.font.size = Pt(9.5)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_box(title, content):
    table = report.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.2)
    hdr = table.cell(0, 0)
    body = table.cell(1, 0)
    shade_cell(hdr, 'B7C9E2')
    shade_cell(body, 'F7F7F7')
    set_cell_text(hdr, title, bold=True)
    p = body.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(content)
    run.font.name = 'Courier New'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
    run.font.size = Pt(9)
    report.add_paragraph('')


def add_heading(text, level=1):
    report.add_heading(text, level=level)


def add_para(text, bold=False, center=False):
    p = report.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
    run.font.size = Pt(10)
    run.bold = bold
    return p


def add_bullets(items):
    for item in items:
        p = report.add_paragraph(style='Normal')
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run('- ' + item)
        run.font.name = 'Courier New'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
        run.font.size = Pt(10)


def add_code(text):
    p = report.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
    run.font.size = Pt(9)


def add_table(headers, rows, col_widths=None):
    table = report.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, bold=True, align='center')
        shade_cell(hdr_cells[i], 'D9EAD3')
        if col_widths:
            hdr_cells[i].width = col_widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            align = 'center' if len(str(val)) <= 10 else 'left'
            set_cell_text(cells[i], str(val), align=align)
            if col_widths:
                cells[i].width = col_widths[i]
    report.add_paragraph('')

# Cover
add_para('PicoRV32 Islemcisi Icin RV32I Alt Kumesi Assembler Tasarimi ve Gerceklestirimi', bold=True, center=True)
add_para('C99 ile Iki Gecisli (Two-Pass) Assembler ve Intel HEX Uretimi', center=True)
add_para('Hazirlayan: [Ogrenci adi soyadi buraya yazilabilir]', center=True)
add_para('Ders: Sistem Programlama / Bilgisayar Mimarisi', center=True)
add_para('Rapor dili: Turkce   Font: Courier New 10 pt', center=True)
report.add_page_break()

add_heading('1. Giris', 1)
add_para('Assembler, assembly kaynak kodunu makine tarafindan dogrudan yorumlanabilir nesne koda ceviren sistem yazilimidir. RISC-V ailesinin acik ve moduler yapiya sahip olmasi, egitsel ve deneysel assembler tasarimlarini ozellikle anlamli hale getirmektedir. Bu projede PicoRV32 cekirdegi hedeflenmis, fakat tasarim karari olarak yalnizca RV32I taban komut kumesinin pratikte gerekli olan alt kumesi uygulanmistir. Amac, label cozumleme, direktif isleme, sabit boyutlu 32-bit komut kodlama ve Intel HEX cikti uretimini tek bir butun olarak gerceklestiren moduler bir assembler gelistirmektir.')
add_para('Gerceklestirilen yazilim C99 ile yazilmistir. Assembler iki gecisli calisir: Pass 1 asamasinda adres atama ve sembol tablosu kurma, Pass 2 asamasinda ise kod uretimi ve nesne dosyasi cikisi yapilir. Bu tercih, ileri referans (forward reference) problemini dogrudan ve guvenilir sekilde cozdugu icin secilmistir.')

add_heading('2. Literatur Arastirmasi', 1)
add_para('Literatur incelemesi iki eksende yapilmistir: (i) assembler tasariminin klasik prensipleri, (ii) RISC-V/PicoRV32 baglamina ozgu kodlama gereksinimleri. Grune ve digerlerinin assembler bolumu, lexical analysis, symbol table yonetimi ve backpatching benzeri mekanizmalarin assembler icin de gecerli oldugunu vurgular. Bu nedenle parser, sembol tablosu ve kodlayici bileşenleri birbirinden ayrik moduller olarak ele alinmistir [4].')
add_para('RISC-V literaturu tarafinda Waterman, Lee, Patterson ve Asanovic tarafindan yayinlanan ISA dokumanlari temel tasarim referansidir. Bu kaynaklar, RV32I komut formatlarinin sabit 32-bit oldugunu, register alanlarinin tum formatlarda ayni bit konumlarinda tutuldugunu ve immediate alanlarinin isaret genisletme mantigina gore yerlestirildigini belirtir [1][2]. Bu sabitlik, assembler tarafinda kod uretimini belirgin sekilde kolaylastirmaktadir.')
add_para('Asanovic ve Patterson tarafindan sunulan acik ISA savunusu, RISC-V ekosisteminin arastirma ve egitim odakli projelerde neden hizla benimsendigini ortaya koymaktadir [3]. PicoRV32 de bu ekosistemin cok yaygin, kucuk alanli ve yapilandirilabilir bir cekirdegi olarak one cikmaktadir. Teknik dokumani, cekirdegin RV32I veya RV32IMC gibi farkli konfigürasyonlarda kullanilabildigini ve yazilim araclarinin bu esnekligi dikkate almasi gerektigini gostermektedir [7].')
add_para('Assembler veri yapilari baglaminda Oeztekin, Lazzem ve Pehlivan tarafindan yapilan calisma, OPTAB arama surelerinin sistem verimini dogrudan etkiledigini ve bu nedenle uygun arama/veri yapisi seciminin onemli oldugunu gostermektedir [5]. Bu projede symbol table icin hash tablo secimi bu motivasyonla yapilmistir. Nakano ve Ito ise assembler tasariminin egitsel ortamlarda yalniz basina degil, islemci ve derleyici ile birlikte ogretildiginde daha verimli sonuclar verdigini gostermistir [6]. Bu raporda blok diyagrami, akis diyagrami ve testler bu pedagojik yaklasimla detaylandirilmistir.')
add_bullets([
    'Klasik assembler islevleri: mnemonic -> opcode cevirisi, adres atama, symbol resolution, object code üretimi.',
    'RISC-V ozelligi: tum temel komutlar 32 bittir; bu durum Pass 1 boyut hesabini sadeleştirir.',
    'PicoRV32 baglami: cekirdek yapilandirilabilir oldugu icin assembler tasariminda desteklenen alt kume acikca sinirlanmalidir.',
    'Veri yapisi sonucu: SYMTAB icin ortalama O(1) arama hedeflenmistir; OPTAB ise kucuk ve sabit oldugu icin salt-okunur tablo olarak tutulmustur.'
])

add_heading('3. Assembler Mimarisi', 1)
add_para('Yazilim mimarisi yedi ana modülden olusur: Parser, LineVector, Opcode Table, Symbol Table, Pass 1, Pass 2, Memory Image/Intel HEX Writer. Parser kaynak satiri ayrıştırır; Pass 1 adres atar ve label kaydeder; Pass 2 ise komutlari bit alanlarina yerlestirerek 32-bit makine kodu uretir. Bellek imgesi dolduktan sonra Intel HEX kayitlari uretilir.')
add_box('Blok Diyagram', r'''
+--------------------+
| Assembly Source .s |
+---------+----------+
          |
          v
+--------------------+
| Parser / Line Split|
+---------+----------+
          |
          v
+--------------------+        +--------------------+
| Pass 1             |------->| Symbol Table       |
| Address Assignment |        | (hash, label->addr)|
+---------+----------+        +--------------------+
          |
          v
+--------------------+        +--------------------+
| Pass 2             |------->| Opcode Table       |
| Encode Instructions|        | (static RV32I map) |
+---------+----------+        +--------------------+
          |
          v
+--------------------+
| Memory Image       |
+---------+----------+
          |
          v
+--------------------+
| Intel HEX Writer   |
+--------------------+
''')
add_para('Mimari, sinirli ama acikca tanimlanmis sorumluluklara sahip moduller uzerine kuruludur. Bu sekilde parser, kodlayici ve cikti uretimi birbirinden bagimsiz test edilebilir. Ayrica listing (.lst) dosyasi sayesinde dogrulama asamasi yalniz HEX cikisina bagli birakilmamistir.')

add_heading('4. Kullanilan Veri Yapilari', 1)
add_para('Projede secilen veri yapilari dogrudan proje gereksinimlerine gore belirlenmistir. Kaynak satirlar LineVector adli dinamik dizi uzerinde tutulur. Bu yapi, iki gecis arasinda ayni ara temsili yeniden kullanmaya imkan verir. Pass 1 sonunda line nesnelerine hesaplanan adresler yazilir; Pass 2 bu bilgiler uzerinden tekrar dosya ayrıştırmak zorunda kalmaz.')
add_para('Symbol Table acik adreslemeli hash tablo olarak uygulanmistir. Her sembol icin isim, adres, section, absolute/relative bayragi ve tanimlanma durumu tutulur. Bu tasarim, ileri referanslarin cozulmesinde ortalama sabit zamanli arama saglar. Opcode Table ise sabit ve kucuk oldugu icin compile-time statik dizi olarak tutulmustur. Bellek imgesi iki paralel dizi ile tutulur: data[] gercek byte degerlerini, used[] ise ilgili adresin ciktiya dahil edilip edilmeyecegini belirtir.')
add_table(
    ['Yapi', 'Amac', 'Gerceklestirim', 'Neden Secildi'],
    [
        ['LineVector', 'Ara temsil', 'Dinamik dizi', 'Pass1/Pass2 arasinda tekrar ayrıştırma gereksinimini kaldirir'],
        ['Symbol Table', 'Label -> adres', 'Hash tablo', 'Ortalama O(1) arama ve ekleme'],
        ['Opcode Table', 'Mnemonic -> kod alanlari', 'Sabit dizi', 'Kucuk ve degismeyen veri kumesi'],
        ['Memory Image', 'Adres -> byte', 'Paralel data/used dizileri', 'Intel HEX üretimi icin dogrudan haritalama'],
        ['ErrorList', 'Hata toplama', 'Sabit sinirli dizi', 'Hatayi tek satirda degil tum derleme akisinda raporlamak']
    ],
    [Cm(3.2), Cm(3.2), Cm(3.2), Cm(6.0)]
)

add_heading('5. Opcode Table Tasarimi', 1)
add_para('Opcode tablosu her komut icin mnemonic, format, operand kalibi, opcode, funct3 ve gerekiyorsa funct7 alanlarini tutar. RISC-V taban komutlari icin opcode ve sabit fonksiyon alanlari assembler tarafinda dogrudan tabloda saklanabilir. Boylece Pass 2 asamasinda mnemonic e gore ilgili tablo girdisi bulunur ve register/immediate alanlari uygun bit konumlarina yerlestirilir.')
add_para('Bu projede desteklenen alt kume; R-type aritmetik/mantiksal komutlar, I-type immediate ve load komutlari, S-type store komutlari, B-type kosullu dallanmalar, U-type ust immediate komutlari, J-type jal komutu ile ecall ve ebreak sistem komutlarini kapsar. FENCE ve CSR ailesi kapsam disi tutulmustur; cunku hedef PicoRV32/RV32I egitsel alt kumesinde zorunlu olmayan ve bu projenin ana hedefi olan temel assembler akisini etkilemeyen unsurlerdir.')
add_table(['Mnemonic','Fmt','Opcode','funct3','funct7','Operand Kalibi'], opcodes, [Cm(2.2), Cm(1.1), Cm(1.8), Cm(1.6), Cm(1.8), Cm(6.0)])

add_heading('6. Symbol Table Tasarimi', 1)
add_para('SYMTAB her label icin benzersiz anahtar degeri gorevini gorur. Pass 1 esnasinda label goruldugunde mevcut LOCCTR degeri ile tabloya kaydedilir. Eger ayni label ikinci kez tanimlanirsa derleme hatasi olusturulur. Pass 2 asamasinda branch/jump ve data referanslari bu tablo uzerinden cozulur.')
add_para('Hash tablonun acik adreslemeli secilmesinin iki nedeni vardir. Birincisi, C99 ile fazladan bagimli kutuphaneler kullanmadan kompakt ve hizli bir uygulama elde etmektir. Ikincisi, label isimlerinin genellikle benzer on ekler icermesi durumunda bile FNV-1a tabanli hash kullanimi ile iyi dagilim saglamaktir. Yuk faktoru yukselince tablo iki katina cikarilir ve rehash islemi yapilir.')
add_bullets([
    'Anahtar: label ismi',
    'Deger: adres, section, absolute/relative bilgisi, tanimli-mi bayragi',
    'Hata durumlari: duplicate label, undefined symbol, gecersiz ifade',
    'Kullanildigi yerler: branch hedefleri, jal/jalr hedefleri, .word/.byte ifadeleri ve .org cozumlemesi'
])

add_heading('7. Assembler Algoritmasi', 1)
add_para('Algoritma iki gecislidir. Sabit 32-bit komut uzunlugu sayesinde instruction size hesabı basitlestirilmis, fakat direktifler ve .org sayesinde adres atama asamasi yine de ayrik bir Pass 1 gerektirmistir. Parser once satiri label, mnemonic ve operand alanlarina ayirir. Ardindan Pass 1 ve Pass 2 ayni satirlar uzerinde ilerler.')
add_box('Akis Diyagrami', r'''
BASLA
  |
  v
Kaynak dosyayi oku
  |
  v
Satirlari ayrıştır
  |
  v
PASS 1:
  - label varsa SYMTAB'a ekle
  - direktif ise LOCCTR guncelle
  - komut ise +4 byte ilerle
  |
  v
Hata var mi?
  |---- Evet ---> Hata listesi yaz ve DUR
  |
  Hayir
  v
PASS 2:
  - direktif verisini uret
  - mnemonic icin OPTAB kaydi bul
  - register/immediate alanlarini kodla
  - Memory Image'a yaz
  |
  v
Intel HEX ve listing uret
  |
  v
BITIR
''')
add_para('Pass 1 islevleri: (i) adres atama, (ii) SYMTAB olusturma, (iii) direktif boyutlarini hesaplama. Pass 2 islevleri: (i) instruction encoding, (ii) veri direktiflerini byte dizisine cevirme, (iii) Memory Image olusturma, (iv) Intel HEX yazma. Parser sonuclarini LineVector icinde saklamak, iki gecis arasinda ara dosya mantigini bellek ici bir temsil ile saglamistir.')

add_heading('8. Test Senaryolari ve Sonuclar', 1)
add_para('Dogrulama icin uc ayri test programi kullanilmistir. Testler bilincli olarak farkli format ailelerini ve direktif kombinasyonlarini icerecek sekilde secilmistir. Her test sonucunda hem Intel HEX hem de listing dosyasi uretilmistir.')
add_table(
    ['Test', 'Amac', 'Kapsam', 'Sonuc'],
    [
        ['Test-1', 'Dongu/sayac', 'addi, blt, ecall', 'Basarili'],
        ['Test-2', 'Veri alani erisimi', '.data, .org, lw, sw, jal, .word', 'Basarili'],
        ['Test-3', 'Alt program cagrisi', 'jal, jalr, add, ecall', 'Basarili']
    ],
    [Cm(2.0), Cm(4.0), Cm(7.0), Cm(2.5)]
)
add_para('Test-1 kaynak programi, geriye dallanan bir dongu ile B-type kodlamasini dogrulamaktadir. Listing dosyasinda blt komutunun FE ile biten negatif ofsetli kod ureterek loop etiketine dondugu gozlenmistir.')
add_code('00000000    93 02 00 00    addi t0, zero, 0')
add_code('00000004    13 03 50 00    addi t1, zero, 5')
add_code('00000008    93 82 12 00    addi t0, t0, 1')
add_code('0000000C    E3 CE 62 FE    blt  t0, t1, loop')
add_code('00000010    73 00 00 00    ecall')
add_para('Test-2, kod ve veri alanlarini farkli adreslere tasimak icin .org ile birlikte .data/.text gecislerini kullanir. Boylece assemblerin duz (flat) adres uzayinda sekmeler arasi yerlesim yapabildigi ve Intel HEX formatinda bosluklari atlayarak yalniz kullanilan adresleri yazdigi dogrulanmistir.')
add_code('00000008    03 A3 02 00    lw   t1, 0(t0)')
add_code('00000010    23 A2 62 00    sw   t1, 4(t0)')
add_code('00000100    0A 00 00 00 14 00 00 00    .word 10, 20')
add_code('00000108    73 00 10 00    ebreak')
add_para('Test-3, alt program cagrisi ve donusu uzerinden J-type ve I-type JALR kodlamasini sinar. JAL ofset hesabinin dogru yapildigi ve ra yazmacina donus adresi yazildigi listing kodundan da gorulebilmektedir.')
add_code('00000008    EF 00 80 00    jal  ra, add_func')
add_code('00000010    33 05 B5 00    add  a0, a0, a1')
add_code('00000014    67 80 00 00    jalr zero, 0(ra)')
add_para('Uygulama duzeyinde bakildiginda testler; ileri label cozumleme, negatif branch ofseti, veri yazimi, mutlak adresli .org, ve alt program akislarini kapsadigi icin proje gereksinimlerinin cekirdek kismi yeterli bicimde dogrulanmistir.')

add_heading('9. Algoritma Karmasikligi Analizi', 1)
add_para('N, kaynak satir sayisini; S, sembol sayisini; B, uretilen toplam byte sayisini gostersin. OPTAB sabit boyutlu oldugu icin arama maliyeti pratikte O(1) kabul edilmistir. SYMTAB icin acik adreslemeli hash tablo kullanildigindan ekleme ve arama ortalama O(1), en kotu durumda ise O(S) olabilir.')
add_table(
    ['Asama', 'Ortalama Zaman', 'En Kotu Zaman', 'Alan'],
    [
        ['Parser', 'O(N)', 'O(N)', 'O(N)'],
        ['Pass 1', 'O(N)', 'O(N*S) patolojik carpismada', 'O(S)'],
        ['Pass 2', 'O(N + B)', 'O(N*S + B)', 'O(B + S)'],
        ['Intel HEX yazimi', 'O(B)', 'O(B)', 'O(1) ek alan']
    ],
    [Cm(3.0), Cm(4.0), Cm(5.5), Cm(3.0)]
)
add_para('Dolayisiyla assemblerin ortalama toplam zamani O(N + B), alan karmasikligi ise O(N + S + B) olarak ifade edilebilir. Cunku ara temsil, sembol tablosu ve bellek imgesi ayni anda tutulmaktadir. Sabit 32-bit komut uzunlugu, Pass 1 tarafinda instruction-size kararini sabit zamana indirdigi icin RV32I baglaminda tasarimi belirgin sekilde sadeleştirmistir.')

add_heading('10. Program Ciktilari', 1)
add_para('Not: PÇ tanimlari kurum ve bolum bazinda farklilik gosterebildigi icin bu bolumde PÇ6, PÇ7, PÇ8, PÇ12 ve PÇ13 degerlendirmesi, tanimlar acik verilmediginden proje urununun olculebilir katkisi uzerinden yapilmistir.')
add_heading('10.1 PÇ6 Kapsaminda Katki', 2)
add_para('Proje, sistem yazilimi geliştirme surecinde problem cozumune yonelik analitik dusunmeyi ve soyutlamadan gercek bit duzeyine inebilmeyi gostermektedir. Instruction formatlarinin tek tek alanlara ayrilmasi, branch/jump immediate yeniden duzenleme mantiginin kurulmasi ve nesne kodun dogrudan uretilebilmesi PÇ6 kapsaminda somut bir urundur.')
add_heading('10.2 PÇ7 Kapsaminda Katki', 2)
add_para('Moduler C99 tasarimi, kaynak kod organizasyonu, derleme otomasyonu ve tekrar uretilebilir testler yazilim muhendisligi disiplinini desteklemektedir. Parser, encoder ve writer modullerinin ayrilmasi; bakim, genisletme ve birim test gelistirme acisindan olumlu bir yazilim mimarisi sergilemektedir.')
add_heading('10.3 PÇ8 Kapsaminda Katki', 2)
add_para('Test programlari ve listing dosyalari, sadece programin calismasini degil ayni zamanda uretilen bit desenlerinin dogrulanmasini da mumkun kilmistir. Bu, deney tasarimi, sonuc yorumlama ve dogrulama becerileri acisindan anlamli bir cikti sunmaktadir.')
add_heading('10.4 PÇ12 Kapsaminda Katki', 2)
add_para('RISC-V ve PicoRV32 gibi acik ekosistemlerin kullanilmasi, guncel teknolojileri takip etme ve yeni mimarilere uyarlanabilir sistem araci gelistirme yetkinligini gostermektedir. Intel HEX cikisi secimi de aracin FPGA/prototipleme akislariyla uyumlu olmasini saglamistir.')
add_heading('10.5 PÇ13 Kapsaminda Katki', 2)
add_para('Raporlama, kaynak kod belgelendirme, listing uretimi ve testlerin acik sunumu proje sonucunun yalnizca kod degil, teknik iletisim urunu olarak da olgunlastigini gostermektedir. Bu durum sunum ve teknik dokumantasyon boyutunu guclendirmektedir.')

add_heading('11. Sonuc ve Degerlendirme', 1)
add_para('Bu projede PicoRV32 odakli RV32I alt kumesi icin tam calisan, iki gecisli ve Intel HEX cikisi ureten bir assembler gelistirilmistir. Proje; opcode tablosu, sembol tablosu, parser, kodlayici ve HEX yazici modullerinden olusan temiz bir mimariye sahiptir. Uc test programi ile branch, load/store, veri yerlesimi, alt program cagrisi ve direktif isleme mekanizmalari dogrulanmistir.')
add_para('Uygulamanin guclu yanlari; sabit 32-bit komut ailesine uygun sade tasarim, moduler kaynak kod, listing dosyasi ile gozlenebilirlik, ve hash tabanli sembol tablosudur. Sinirlari ise; pseudo-instruction destegi, FENCE/CSR ailesi, relocation record uretimi, macro desteği ve daha zengin ifade cozumleyicisinin henuz eklenmemis olmasidir. Bunlar gelecekteki dogal gelistirme maddeleri olarak gorulmektedir.')
add_para('Sonuc olarak yazilim, ders kapsami icin yeterli derinlikte bir assembler altyapisi sunmakta; mimari tasarim, veri yapilari, algoritma, yazilim gerceklestirme ve test/dogrulama basliklarinda yuksek puan bandini hedefleyen tutarli bir proje urunu ortaya koymaktadir.')

add_heading('12. Kaynakca', 1)
refs = [
    '[1] A. Waterman, Y. Lee, D. A. Patterson, K. Asanovic, The RISC-V Instruction Set Manual, Volume I: User-Level ISA, Version 2.0, UCB/EECS-2014-54, University of California, Berkeley, 2014.',
    '[2] A. Waterman, K. Asanovic (eds.), The RISC-V Instruction Set Manual, Volume I: Unprivileged Architecture, RISC-V Ratified Specification Library, 2025.',
    '[3] K. Asanovic, D. A. Patterson, Instruction Sets Should Be Free: The Case For RISC-V, UCB/EECS-2014-146, University of California, Berkeley, 2014.',
    '[4] D. Grune, C. J. H. Jacobs, K. G. Langendoen, et al., Modern Compiler Design, 2nd ed., Springer, Chapter 8: Assemblers, Disassemblers, Linkers, and Loaders, 2012.',
    '[5] H. Oztekin, A. Lazzem, I. Pehlivan, Using FPGA-based content-addressable memory for mnemonics instruction searching in assembler design, The Journal of Supercomputing, 79(15):17386-17418, 2023.',
    '[6] K. Nakano, Y. Ito, Processor, Assembler, and Compiler Design Education Using an FPGA, Proceedings of ICPADS 2008, pp. 723-728, IEEE, 2008.',
    '[7] YosysHQ, PicoRV32 README and technical documentation, GitHub repository, accessed 2026.'
]
for ref in refs:
    add_para(ref)

# Save
report.save(OUT)
print(OUT)
